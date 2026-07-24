from __future__ import annotations

"""Build the verified Korean full-stack, training, QA, and E2E Word manual.

The document intentionally separates functional pipeline acceptance from the
independent 95% appearance/manufacturing target.  It is generated from the
current repository defaults and the evidence collected on 2026-07-23.
"""

import argparse
import os
import sys
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
STACK_ROOT = SCRIPT_DIR.parent
REPO_ROOT = STACK_ROOT.parent
DOCS_DIR = STACK_ROOT / "docs"
ASSET_DIR = STACK_ROOT / "storage" / "manual-assets" / "20260723"
DEFAULT_OUTPUT = DOCS_DIR / "Xconcep_AI_전체_파이프라인_운영_학습_QA_E2E_서버이식_매뉴얼_20260723.docx"
UI_EVIDENCE = (
    STACK_ROOT
    / "storage"
    / "e2e-evidence"
    / "20260723-full-regression"
    / "ui-meeting-structured-result.png"
)

DOCUMENTS_SKILL = Path(
    os.environ.get(
        "CODEX_DOCUMENTS_SKILL",
        r"C:\Users\user\.codex\plugins\cache\openai-primary-runtime"
        r"\documents\26.715.12143\skills\documents",
    )
)
sys.path.insert(0, str(DOCUMENTS_SKILL / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "1565C0"
CYAN = "00A6C7"
NAVY = "092033"
INK = "243447"
MUTED = "5B6B7C"
LIGHT_BLUE = "EAF3FB"
LIGHT_CYAN = "E8F8FA"
LIGHT_GRAY = "F2F5F7"
MID_GRAY = "D5DEE5"
DARK_GRAY = "455A64"
GREEN = "1B8A5A"
AMBER = "B56A00"
RED = "B3261E"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, values in edges.items():
        tag = f"w:{edge_name}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in values.items():
            element.set(qn(f"w:{key}"), str(value))


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_run_font(run, name: str = "Calibri", east_asia: str = "Malgun Gothic") -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)


def _set_style_font(style, name: str, size: float, *, bold: bool = False, color: str = INK) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _set_repeat_table_header(row) -> None:
    _repeat_header(row)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)


def _add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((run_color, underline))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\malgun.ttf"),
        Path(r"C:\Windows\Fonts\malgunbd.ttf") if bold else Path(r"C:\Windows\Fonts\malgun.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: Sequence[str],
    *,
    fill: str,
    outline: str,
    title_color: str = "#FFFFFF",
    body_color: str = "#E7F0F7",
) -> None:
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, x2, _ = box
    draw.text((x1 + 24, y1 + 18), title, fill=title_color, font=_font(31, bold=True))
    y = y1 + 69
    for line in body:
        draw.text((x1 + 24, y), line, fill=body_color, font=_font(23))
        y += 34


def _arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    color: str = "#47D7E8",
    width: int = 6,
) -> None:
    draw.line((*start, *end), fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 12), (ex - 18 * direction, ey + 12)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 12, ey - 18 * direction), (ex + 12, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def build_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1120), "#071626")
    draw = ImageDraw.Draw(image)
    draw.text((80, 48), "Xconcep AI — 실행 아키텍처와 책임 경계", fill="#FFFFFF", font=_font(45, bold=True))
    draw.text(
        (82, 108),
        "상태 제어와 생성 실행을 분리하고, 모든 산출물은 공유 storage 계약으로 연결",
        fill="#A8BACB",
        font=_font(25),
    )

    _rounded_box(
        draw,
        (80, 190, 430, 400),
        "사용자 계층",
        ["직접 프롬프트·이미지", "회의 음성 Chunk", "3D Viewer·다운로드"],
        fill="#113854",
        outline="#2DB7D5",
    )
    _rounded_box(
        draw,
        (550, 190, 900, 400),
        "Web / Control",
        ["PHP Reverse Proxy", "Django REST API", "MySQL·Celery 상태"],
        fill="#123C5C",
        outline="#55B7FF",
    )
    _rounded_box(
        draw,
        (1020, 190, 1370, 400),
        "Agent / Knowledge",
        ["결정론 Gateway", "NAT 선택 Profile", "Qdrant RAG"],
        fill="#153B50",
        outline="#51D2C8",
    )
    _rounded_box(
        draw,
        (1390, 190, 1720, 400),
        "AI Worker",
        ["STT·요구 분석", "2D·3D Router", "검증·OpenUSD"],
        fill="#17384A",
        outline="#4DE0B5",
    )
    _arrow(draw, (430, 295), (550, 295))
    _arrow(draw, (900, 295), (1020, 295))
    _arrow(draw, (1370, 295), (1390, 295))

    providers = [
        ((80, 520, 385, 720), "2D", ["기본: ComfyUI/FLUX", "선택: OpenAI Image", "4안 + 품질 검사"], "#12364D", "#31C3E1"),
        ((430, 520, 735, 720), "3D", ["TripoSR 빠른 Mesh", "OpenSCAD 파라메트릭", "Blender 고품질"], "#163548", "#42D6C4"),
        ((780, 520, 1085, 720), "Voice / LLM", ["Faster-Whisper", "NeMo/NIM 선택", "Rules/vLLM 선택"], "#17384A", "#72D4F1"),
        ((1130, 520, 1435, 720), "Data", ["MySQL 8.4", "Redis·Celery", "Qdrant 1.18.2"], "#183748", "#69D3A4"),
        ((1480, 520, 1720, 720), "Native", ["OpenSCAD", "Blender", "OpenUSD pxr"], "#173847", "#6FD9B9"),
    ]
    for box, title, body, fill, outline in providers:
        _rounded_box(draw, box, title, body, fill=fill, outline=outline)
        _arrow(draw, (1555, 400), ((box[0] + box[2]) // 2, box[1]))

    _rounded_box(
        draw,
        (180, 835, 1620, 1020),
        "공유 산출물 경계  ·  ./storage",
        [
            "Project / Meeting / Transcript / DesignState / GeometryContract",
            "Concept PNG · GLB · STL · SCAD · USDA · USDC · Manifest · Validation · Evidence",
        ],
        fill="#0F3045",
        outline="#2EC5D3",
    )
    for x in (230, 580, 930, 1280, 1600):
        _arrow(draw, (x, 720), (x, 835))
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, "PNG", optimize=True)


def build_workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 760), "#F4F8FB")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "E2E 실행 흐름과 품질 판정 분리", fill="#092033", font=_font(44, bold=True))
    draw.text(
        (72, 105),
        "기능 통과는 파일 생성·계약 준수를 뜻하며, 0.95 외관 목표 통과와 동일하지 않음",
        fill="#516577",
        font=_font(25),
    )
    labels = [
        ("1", "입력", "텍스트·이미지\n또는 회의 음성"),
        ("2", "구조화", "STT·요구사항\nDesignSpec"),
        ("3", "2D", "FLUX 기본\n4개 콘셉트"),
        ("4", "선택", "사용자 선택\n또는 추천"),
        ("5", "3D", "TripoSR / OpenSCAD\n/ Blender"),
        ("6", "검증", "계약·파일·치수\nOpenUSD 재개방"),
        ("7", "승인", "엔지니어 검토\n제조 승인 별도"),
    ]
    x = 55
    for index, title, body in labels:
        box = (x, 250, x + 215, 565)
        draw.rounded_rectangle(box, radius=22, fill="#FFFFFF", outline="#8CB7D4", width=3)
        draw.ellipse((x + 67, 275, x + 147, 355), fill="#1261A0")
        num_font = _font(34, bold=True)
        bbox = draw.textbbox((0, 0), index, font=num_font)
        draw.text(
            (x + 107 - (bbox[2] - bbox[0]) / 2, 315 - (bbox[3] - bbox[1]) / 2 - 2),
            index,
            fill="#FFFFFF",
            font=num_font,
        )
        draw.text((x + 27, 375), title, fill="#092033", font=_font(28, bold=True))
        y = 425
        for line in body.split("\n"):
            draw.text((x + 27, y), line, fill="#4E6374", font=_font(21))
            y += 34
        if index != "7":
            _arrow(draw, (x + 215, 410), (x + 245, 410), color="#008EA8", width=5)
        x += 245
    draw.rounded_rectangle((1070, 615, 1710, 705), radius=18, fill="#FFF2DE", outline="#D98B20", width=2)
    draw.text(
        (1100, 642),
        "현재 실제 외관 점수 0.4161 / 0.95  →  기능 PASS, 품질 FAIL",
        fill="#8A4E00",
        font=_font(24, bold=True),
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, "PNG", optimize=True)


def build_training_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 820), "#071626")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "CAD VLM 학습 데이터와 운영 연결", fill="#FFFFFF", font=_font(44, bold=True))
    draw.text(
        (72, 105),
        "모델은 임의 CAD 코드를 실행하지 않고, 검증된 DesignSpec JSON만 생성",
        fill="#A8BACB",
        font=_font(25),
    )
    boxes = [
        ((60, 230, 345, 525), "원본", ["DXF / STEP", "PHP 패키지", "권리·출처"]),
        ((405, 230, 690, 525), "전처리", ["manifest/geometry", "품질 ≥ 0.90", "해시·PNG"]),
        ((750, 230, 1035, 525), "통합 데이터", ["records.jsonl", "split 누수 차단", "DesignSpec 정답"]),
        ((1095, 230, 1380, 525), "학습", ["Qwen3-VL", "LoRA / QLoRA", "체크포인트"]),
        ((1440, 230, 1740, 525), "승격", ["독립 홀드아웃", "Wilson 95% 하한", "E2E 재검증"]),
    ]
    for i, (box, title, body) in enumerate(boxes):
        _rounded_box(draw, box, title, body, fill="#12394F", outline="#3AC1D5")
        if i < len(boxes) - 1:
            _arrow(draw, (box[2], 377), (boxes[i + 1][0][0], 377))
    draw.rounded_rectangle((230, 630, 1570, 755), radius=20, fill="#113044", outline="#50D8B8", width=3)
    draw.text(
        (270, 654),
        "승격 모델 → JSON 스키마 검증 → build_geometry_contract() →",
        fill="#E6F5FA",
        font=_font(25, bold=True),
    )
    draw.text(
        (270, 698),
        "OpenSCAD / Blender / OpenUSD → 제조성·외관 독립 평가",
        fill="#C7E7EF",
        font=_font(25),
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, "PNG", optimize=True)


class ManualBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        styles = self.doc.styles
        normal = styles["Normal"]
        _set_style_font(normal, "Calibri", 11, color=INK)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        normal.paragraph_format.widow_control = True

        for style_name, size, before, after, color in (
            ("Title", 28, 0, 14, NAVY),
            ("Subtitle", 14, 0, 12, MUTED),
            ("Heading 1", 16, 18, 10, BLUE),
            ("Heading 2", 13, 14, 7, NAVY),
            ("Heading 3", 12, 10, 5, DARK_GRAY),
        ):
            style = styles[style_name]
            _set_style_font(style, "Calibri", size, bold=True, color=color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        for style_name in ("List Bullet", "List Number"):
            style = styles[style_name]
            _set_style_font(style, "Calibri", 10.5, color=INK)
            style.paragraph_format.left_indent = Inches(0.375)
            style.paragraph_format.first_line_indent = Inches(-0.188)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.25

        if "Code Block" not in styles:
            code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        else:
            code_style = styles["Code Block"]
        _set_style_font(code_style, "Consolas", 8.5, color="193549")
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.12)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.line_spacing = 1.05
        code_style.paragraph_format.keep_together = True

        if "Caption" in styles:
            caption = styles["Caption"]
            _set_style_font(caption, "Calibri", 9, color=MUTED)
            caption.font.italic = True
            caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(8)

        self._configure_header_footer(section)

    def _configure_header_footer(self, section) -> None:
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.autofit = False
        apply_table_geometry(
            table,
            [6200, 3160],
            table_width_dxa=TABLE_WIDTH_DXA,
            indent_dxa=0,
            cell_margins_dxa={"top": 30, "bottom": 30, "start": 20, "end": 20},
        )
        table.cell(0, 0).text = "X concep AI  ·  기술 운영 매뉴얼"
        table.cell(0, 1).text = "VERIFIED 2026-07-23"
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                _set_run_font(run)
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(MUTED)
            _set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": CYAN})

        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        table.autofit = False
        apply_table_geometry(
            table,
            [7600, 1760],
            table_width_dxa=TABLE_WIDTH_DXA,
            indent_dxa=0,
            cell_margins_dxa={"top": 30, "bottom": 20, "start": 20, "end": 20},
        )
        left = table.cell(0, 0).paragraphs[0]
        left.add_run("내부 기술 검토용 · 제조 승인 문서 아님")
        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.add_run("v1.0  ·  ")
        _page_field(right)
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                _set_run_font(run)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(MUTED)

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def paragraph(self, text: str = "", *, bold_prefix: str | None = None) -> None:
        paragraph = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            run = paragraph.add_run(bold_prefix)
            run.bold = True
            _set_run_font(run)
            remainder = paragraph.add_run(text[len(bold_prefix) :])
            _set_run_font(remainder)
        else:
            run = paragraph.add_run(text)
            _set_run_font(run)

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(item)
            _set_run_font(run)

    def numbers(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph(style="List Number")
            run = paragraph.add_run(item)
            _set_run_font(run)

    def code(self, text: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        apply_table_geometry(
            table,
            [TABLE_WIDTH_DXA],
            table_width_dxa=TABLE_WIDTH_DXA,
            indent_dxa=TABLE_INDENT_DXA,
            cell_margins_dxa={"top": 100, "bottom": 100, "start": 150, "end": 150},
        )
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "F2F6F8")
        _set_cell_border(cell, left={"val": "single", "sz": "18", "color": CYAN})
        paragraph = cell.paragraphs[0]
        paragraph.style = self.doc.styles["Code Block"]
        for index, line in enumerate(text.rstrip().splitlines()):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            _set_run_font(run, "Consolas", "Malgun Gothic")
            run.font.size = Pt(8.5)

    def callout(self, title: str, body: str, *, tone: str = "info") -> None:
        colors = {
            "info": (LIGHT_BLUE, BLUE),
            "success": ("E8F5EE", GREEN),
            "warning": ("FFF3DF", AMBER),
            "danger": ("FDECEA", RED),
        }
        fill, accent = colors[tone]
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        apply_table_geometry(
            table,
            [TABLE_WIDTH_DXA],
            table_width_dxa=TABLE_WIDTH_DXA,
            indent_dxa=TABLE_INDENT_DXA,
            cell_margins_dxa={"top": 110, "bottom": 110, "start": 160, "end": 160},
        )
        cell = table.cell(0, 0)
        _set_cell_shading(cell, fill)
        _set_cell_border(cell, left={"val": "single", "sz": "22", "color": accent})
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(title + "  ")
        _set_run_font(run)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(accent)
        run = paragraph.add_run(body)
        _set_run_font(run)
        run.font.size = Pt(10.5)

    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[object]],
        widths: Sequence[int],
        *,
        font_size: float = 8.6,
    ) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.autofit = False
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.cell(0, index)
            cell.text = str(header)
            _set_cell_shading(cell, "DCE8F1")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for run in cell.paragraphs[0].runs:
                _set_run_font(run)
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor.from_string(NAVY)
        _set_repeat_table_header(table.rows[0])
        _keep_row_together(table.rows[0])
        for row_index, values in enumerate(rows):
            row = table.add_row()
            if row_index % 2:
                for cell in row.cells:
                    _set_cell_shading(cell, "F7FAFC")
            for index, value in enumerate(values):
                cell = row.cells[index]
                cell.text = str(value)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.08
                    for run in paragraph.runs:
                        _set_run_font(run)
                        run.font.size = Pt(font_size)
            _keep_row_together(row)
        apply_table_geometry(
            table,
            widths,
            table_width_dxa=TABLE_WIDTH_DXA,
            indent_dxa=TABLE_INDENT_DXA,
            cell_margins_dxa=CELL_MARGINS,
        )
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def picture(self, path: Path, caption: str, *, width: float = 6.35) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width))
        caption_paragraph = self.doc.add_paragraph(caption, style="Caption")
        for run in caption_paragraph.runs:
            _set_run_font(run)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def cover(self) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(34)
        run = paragraph.add_run("X")
        _set_run_font(run)
        run.font.size = Pt(30)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), BLUE)
        run._r.get_or_add_rPr().append(shading)
        run = paragraph.add_run("  X concep AI")
        _set_run_font(run)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)

        title = self.doc.add_paragraph(style="Title")
        title.paragraph_format.space_before = Pt(44)
        title.paragraph_format.space_after = Pt(16)
        run = title.add_run("전체 파이프라인·운영·학습\nQA·E2E·서버 이식 매뉴얼")
        _set_run_font(run)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)

        subtitle = self.doc.add_paragraph(style="Subtitle")
        run = subtitle.add_run(
            "ComfyUI/FLUX 기본 2D · OpenAI Image 선택 모드 · TripoSR · "
            "OpenSCAD 파라메트릭 · Blender · Meeting Voice · OpenUSD"
        )
        _set_run_font(run)
        run.font.color.rgb = RGBColor.from_string(MUTED)

        self.doc.add_paragraph()
        self.callout(
            "검증본",
            "2026-07-23 로컬 통합 환경에서 코드 회귀 100건, CAD 학습 경로 15건, "
            "한국어 설비 회의 음성 E2E 및 실제 UI를 확인한 상태를 기준으로 작성함.",
            tone="success",
        )

        table = self.doc.add_table(rows=5, cols=2)
        table.autofit = False
        rows = [
            ("문서 버전", "1.0"),
            ("기준 날짜", "2026-07-23 (Asia/Seoul)"),
            ("대상", "개발·ML·QA·인프라·설비 엔지니어"),
            ("현재 인증", "내부 Docker MySQL / 외부 사내 MySQL은 추후 전환"),
            ("기밀 등급", "내부 기술 검토용"),
        ]
        for row_index, (key, value) in enumerate(rows):
            table.cell(row_index, 0).text = key
            table.cell(row_index, 1).text = value
            _set_cell_shading(table.cell(row_index, 0), LIGHT_GRAY)
            for run in table.cell(row_index, 0).paragraphs[0].runs:
                _set_run_font(run)
                run.font.bold = True
                run.font.size = Pt(9)
            for run in table.cell(row_index, 1).paragraphs[0].runs:
                _set_run_font(run)
                run.font.size = Pt(9)
            _keep_row_together(table.rows[row_index])
        apply_table_geometry(
            table,
            [2100, 7260],
            table_width_dxa=TABLE_WIDTH_DXA,
            indent_dxa=TABLE_INDENT_DXA,
            cell_margins_dxa=CELL_MARGINS,
        )

        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(30)
        run = paragraph.add_run(
            "중요: 이 문서는 AI 생성 결과를 제조 승인으로 간주하지 않습니다. "
            "95%는 통계적으로 검증할 내부 acceptance 목표이며 현재 달성 선언이 아닙니다."
        )
        _set_run_font(run)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(RED)
        self.page_break()

    def build(self) -> None:
        architecture = ASSET_DIR / "architecture.png"
        workflow = ASSET_DIR / "e2e-workflow.png"
        training = ASSET_DIR / "training-flow.png"
        build_architecture_diagram(architecture)
        build_workflow_diagram(workflow)
        build_training_diagram(training)

        self.cover()
        self._document_control()
        self._status_summary()
        self._architecture(architecture)
        self._workflows(workflow)
        self._routing_and_contracts()
        self._installation_and_operation()
        self._configuration()
        self._server_migration()
        self._libraries_and_paths()
        self._training(training)
        self._qa_e2e()
        self._troubleshooting()
        self._cli_reference()
        self._design_rationale()
        self._checklists_and_sources()

    def _document_control(self) -> None:
        self.heading("문서 사용법과 변경 통제", 1)
        self.paragraph(
            "이 문서는 현재 저장소를 기준으로 시스템 전체를 한 번에 설치·운영·학습·검증하기 위한 "
            "기준 매뉴얼이다. 기존의 개별 Markdown 문서는 세부 설계 이력을 보존하며, 상충할 때는 "
            "현재 코드, .env.example, docker-compose.yml, 이 문서 순서로 확인한다."
        )
        self.heading("독자별 빠른 경로", 2)
        self.table(
            ["독자", "먼저 읽을 절", "완료 기준"],
            [
                ("운영/인프라", "6 설치·운영 → 8 서버 이식 → 13 장애 대응", "상태 API·백업·복구·롤백 확인"),
                ("ML 엔지니어", "10 학습 → 11 QA/E2E", "데이터 검증·학습 스모크·독립 홀드아웃"),
                ("QA", "2 현황 → 11 QA/E2E → 15 체크리스트", "기능 PASS와 품질 목표를 분리 판정"),
                ("설비/CAD", "4 워크플로 → 5 계약·등급 → 10.9 승격", "치수·구성·관계 확인 후 엔지니어 검토"),
                ("개발", "3 아키텍처 → 7 설정 → 12 CLI", "서비스 책임 경계와 재현 명령 이해"),
            ],
            [1500, 4300, 3560],
            font_size=8.8,
        )
        self.heading("문서 상태 표현", 2)
        self.bullets(
            [
                "PASS: 실제 실행 또는 자동 테스트에서 판정 조건을 충족함.",
                "FAIL: 기대 계약을 충족하지 못했으며 운영 승격에 사용할 수 없음.",
                "NOT RUN: 연결 또는 비용·권한이 필요한 선택 기능으로 이번 검증에서 실행하지 않음.",
                "STRUCTURED: 구조·배치 검토용이며 완성 CAD나 제조 승인을 의미하지 않음.",
            ]
        )
        self.callout(
            "비밀값 처리",
            "OPENAI_API_KEY, DB 비밀번호, HF_TOKEN, VLM_API_KEY는 이 문서·Git·스크린샷에 기록하지 않는다. "
            "대상 서버의 git-ignored .env 또는 사내 Secret Manager에만 저장한다.",
            tone="warning",
        )

    def _status_summary(self) -> None:
        self.page_break()
        self.heading("1. 현재 검증 상태 요약", 1)
        self.paragraph(
            "최종 검증은 내부 MySQL 인증, 로컬 ComfyUI/FLUX, Faster-Whisper, 네이티브 OpenSCAD·Blender, "
            "TripoSR, OpenUSD가 연결된 환경에서 수행했다. OpenAI Image는 추가 선택 모드로 유지하며 "
            "최종 회의 E2E의 기준 2D 생성기는 ComfyUI였다."
        )
        self.table(
            ["영역", "결과", "증거/수치", "판정"],
            [
                ("서비스 회귀", "DRF 12 + Worker 85 + Agent 1 + Knowledge 2", "총 100 tests, PHP/JS 구문 통과", "PASS"),
                ("CAD 학습 패키지", "전처리·평가·PHP importer 테스트", "15 tests", "PASS"),
                ("회의 음성 E2E", "3개 한국어 설비 회의 MP3", "핵심어 11/12 = 91.67%", "PASS"),
                ("치수 복원", "폭·깊이·높이", "1200 / 800 / 1600 mm", "PASS"),
                ("2D 생성", "ComfyUI/FLUX 4안", "각 1024×1024, 기본 품질 검사 통과", "PASS"),
                ("3D 구조 생성", "openscad_equipment", "GLB 54,884B · STL 594,635B", "PASS"),
                ("OpenUSD", "USDA·USDC·Layered package", "Stage·manifest 산출", "PASS"),
                ("브라우저 UI", "이력·결과·뷰어·ISO/FRONT/TOP·다운로드 표시", "실제 UI 캡처", "PASS"),
                ("외관 자가평가", "선택 2D ↔ 구조 3D", "0.4161 / 목표 0.95", "FAIL"),
                ("학습 실행 가능성", "Qwen3-VL 4B LoRA 1 step", "loss 3.2974 · 39.17s", "PASS(스모크)"),
            ],
            [1900, 3000, 2960, 1500],
            font_size=8.3,
        )
        self.callout(
            "핵심 결론",
            "워크플로의 기능 연결은 PASS다. 그러나 현재 구조화 3D가 선택된 사실적 2D 외관을 95%로 "
            "재현한다는 근거는 없으며 실제 측정도 FAIL이다. 운영 UI의 ‘구조 검토 가능’ 등급이 정확한 표현이다.",
            tone="danger",
        )
        self.heading("1.1 실제 음성 E2E 식별자", 2)
        self.table(
            ["항목", "값"],
            [
                ("Project ID", "PRJ-E76ED59D94"),
                ("Speech", "faster-whisper local / small"),
                ("2D", "ComfyUI · flux-2-klein-base-4b-fp8.safetensors · 4 concepts"),
                ("3D", "openscad_equipment · validation_grade=structured"),
                ("산출물", "GLB, STL, SCAD, Geometry JSON, PNG, USDA, USDC, OpenUSD package/manifest"),
                ("스크린샷", "storage/e2e-evidence/20260723-full-regression/ui-meeting-structured-result.png"),
            ],
            [2300, 7060],
            font_size=8.8,
        )
        if UI_EVIDENCE.exists():
            self.picture(
                UI_EVIDENCE,
                "그림 1. 실제 한국어 회의 음성에서 생성한 설비 파라메트릭 3D 결과 UI",
                width=6.35,
            )

    def _architecture(self, architecture: Path) -> None:
        self.page_break()
        self.heading("2. 전체 시스템 아키텍처", 1)
        self.picture(architecture, "그림 2. 실행 계층·AI 제공자·공유 산출물 경계", width=6.4)
        self.heading("2.1 서비스별 책임", 2)
        self.table(
            ["서비스", "책임", "주요 데이터", "기본 포트/경로"],
            [
                ("frontend-php", "UI, 인증 세션, DRF reverse proxy", "입력·파일·결과 화면", "WEB_PORT→80"),
                ("control-plane-drf", "Project/Job/Meeting 상태의 단일 소유자", "MySQL 레코드·Celery job", "8000 internal"),
                ("celery-worker", "비동기 API 작업 실행", "Redis broker/result", "no host port"),
                ("agent-layer", "고정 workflow와 선택 NAT 경계", "생성·회의 request/response", "8010 internal"),
                ("python-worker", "STT·분석·2D·3D·검증·OpenUSD", "공유 storage artifacts", "8001 internal"),
                ("knowledge-service", "문서 추출·검색·Qdrant adapter", "chunks·vectors·asset index", "8020 internal"),
                ("mysql", "프로젝트와 현재 내부 인증", "Django schema", "127.0.0.1:3307 기본"),
                ("redis", "Celery broker·result", "queue state", "6379 internal"),
                ("qdrant", "과거 프로젝트/RAG memory", "vectors·payload", "6333 internal"),
                ("nat-runtime", "NVIDIA NAT 선택 profile", "tool trace", "8011 / profile=nat"),
            ],
            [1800, 2900, 2660, 2000],
            font_size=8.0,
        )
        self.heading("2.2 책임 경계가 필요한 이유", 2)
        self.bullets(
            [
                "Worker가 MySQL을 직접 수정하지 않게 하여 재시도·감사·권한 정책을 DRF 한 곳에서 통제한다.",
                "AI 모델 교체가 Project/Job 상태 모델을 흔들지 않도록 Gateway와 Worker를 분리한다.",
                "모든 서비스가 동일 ./storage를 마운트해 대형 GLB·음성·이미지를 DB에 넣지 않는다.",
                "NAT, NIM, 외부 MySQL, OpenAI는 선택 Profile로 남겨 로컬 기본 경로를 차단하지 않는다.",
            ]
        )
        self.heading("2.3 저장 데이터", 2)
        self.table(
            ["저장소", "저장 항목", "백업 단위", "복구 확인"],
            [
                ("MySQL", "프로젝트·Job·Meeting·사용자·품질 증거", "논리 dump", "테이블·row count"),
                ("storage/", "업로드·전사·2D·3D·USD·보고서", "파일 스냅샷/rsync", "SHA-256·파일 재개방"),
                ("Qdrant", "검색 벡터·metadata", "Qdrant snapshot", "collection count·search smoke"),
                ("Redis", "실행 중 queue/result", "영구 운영이면 AOF", "queue drain 후 재기동"),
                ("모델 캐시", "Whisper·TripoSR·HF base model", "버전·revision 재다운로드 또는 별도 캐시", "모델 SHA·health"),
            ],
            [1600, 3120, 2300, 2340],
            font_size=8.4,
        )

    def _workflows(self, workflow: Path) -> None:
        self.page_break()
        self.heading("3. 사용자 워크플로", 1)
        self.picture(workflow, "그림 3. 입력부터 엔지니어 승인까지의 E2E 경계", width=6.4)
        self.heading("3.1 직접 입력 워크플로", 2)
        self.numbers(
            [
                "사용자가 설비·모듈·부품 설명과 선택 참고 이미지를 입력한다.",
                "DRF가 Project를 만들고 Worker가 요구사항을 DesignSpec으로 구조화한다.",
                "기본 ComfyUI/FLUX가 서로 다른 방향의 2D 콘셉트 4개를 생성한다.",
                "사용자가 원하는 2D를 선택하거나 추천안을 사용한다.",
                "Router가 요구 유형·복잡도·신뢰도에 따라 TripoSR, 범용 OpenSCAD, 전문 OpenSCAD, Blender 또는 Hybrid를 고른다.",
                "산출물을 재개방하고 치수·필수 구성·관계·단위·OpenUSD 계약을 검증한다.",
                "UI에서 GLB를 검토하고 STL·SCAD·Geometry JSON·USD를 내려받는다.",
            ]
        )
        self.heading("3.2 회의 음성 워크플로", 2)
        self.numbers(
            [
                "브라우저 MediaRecorder 또는 QA fixture가 약 15초 단위 음성 Chunk를 업로드한다.",
                "Faster-Whisper/NeMo/NIM 중 선택된 STT가 Transcript segment를 누적한다.",
                "회의 분석기가 확정 요구, 치수, 부품, 안전, 변경, 미확정 항목을 구조화한다.",
                "사용자가 Transcript와 분석 결과를 확인·수정한다.",
                "분석 결과로 2D 4안을 만들고 선택한 안으로 동일 3D 파이프라인을 실행한다.",
                "후속 발언은 Revision Patch와 OpenUSD revision layer로 남긴다.",
            ]
        )
        self.callout(
            "음성 QA 기준",
            "현재 smoke-meeting-live.ps1은 핵심어 recall 0.75 이상, 폭·깊이·높이 모두 복원, "
            "2D 4안, 구조화 GLB의 유효성까지 요구한다. 최종 실측은 0.9167이었다.",
            tone="success",
        )
        self.heading("3.3 결과 재생성", 2)
        self.table(
            ["UI 선택", "의도", "주 경로", "적합한 검토"],
            [
                ("더 빠르게 생성", "외형을 빨리 확인", "TripoSR", "컨셉·실루엣"),
                ("구조를 정확하게 재생성", "치수·구성·배치 우선", "전문 OpenSCAD", "구조·후속 CAD 입력"),
                ("더 사실적으로 재생성", "재질·조명·조립 시각화", "Blender/Hybrid", "프레젠테이션·외관"),
                ("동작·OpenUSD로 확장", "레이어·variant·physics", "Blender + OpenUSD", "Omniverse/시뮬레이션"),
                ("실패 항목만 재생성", "계약 실패 그룹 한정", "부분 GeometryContract", "재작업 비용 절감"),
            ],
            [2200, 2600, 2100, 2460],
            font_size=8.6,
        )

    def _routing_and_contracts(self) -> None:
        self.page_break()
        self.heading("4. 생성 모드와 라우팅 정책", 1)
        self.heading("4.1 2D 공급자 정책", 2)
        self.table(
            ["모드", "설정", "역할", "현재 상태"],
            [
                ("로컬 기본", "OPENAI_IMAGE_MODE=comfyui", "FLUX 기반 4안·정밀 프롬프트·기본 품질 검사", "실제 E2E PASS"),
                ("추가 모드", "OPENAI_IMAGE_MODE=openai", "기존 프롬프트 제한과 파이프라인 재현성 비교", "선택/비용 제한"),
                ("테스트", "OPENAI_IMAGE_MODE=mock", "GPU/API 없이 UI·상태 배선 확인", "개발 전용"),
            ],
            [1600, 2600, 3360, 1800],
            font_size=8.7,
        )
        self.paragraph(
            "OpenAI Image API Key는 브라우저나 PHP에 노출하지 않고 Worker의 .env에서만 읽는다. "
            "일일 요청수·예상 비용 상한을 넘으면 호출을 차단한다. 이번 최종 회의 E2E는 기본 로컬 모드로 수행했다."
        )
        self.heading("4.2 3D 생성 경로", 2)
        self.table(
            ["route", "출력 성격", "장점", "제한"],
            [
                ("triposr", "선택 2D의 빠른 단일-view mesh", "빠른 웹 미리보기", "치수·조립·후면 구조 정확도 낮음"),
                ("openscad", "기존 범용 template", "fallback 유지·결정론", "형상 범위 제한"),
                ("openscad_auto", "요구에 맞춘 전문 mode 추천", "일관된 route 선택", "잘못된 분류 시 재선택 필요"),
                ("openscad_part", "부품 파라메트릭", "홀·리브·브래킷 치수 계약", "자유 곡면에 부적합"),
                ("openscad_module", "구동/작업 모듈", "부품 관계·조립 구조", "표준부품 라이브러리 확장 필요"),
                ("openscad_equipment", "설비 프레임·컨베이어·안전·제어", "구조·배치·다중뷰 검증", "사실적 외관은 Blender 필요"),
                ("blender", "재질·조명·camera·assembly", "외관·OpenUSD bridge", "생성 시간·script 검증 필요"),
                ("hybrid", "구조 + 사실적 후처리", "구조와 외관 결합", "가장 높은 계산 비용"),
            ],
            [2000, 2500, 2440, 2420],
            font_size=8.0,
        )
        self.heading("4.3 Router 임계값", 2)
        self.table(
            ["설정", "기본값", "의미", "조정 원칙"],
            [
                ("ROUTING_LOW_CONFIDENCE", "0.60", "이하이면 안전 fallback", "실패가 많으면 상향"),
                ("ROUTING_HIGH_CONFIDENCE", "0.80", "이상이면 전문 route", "전문 생성기 정확도 검증 후 하향"),
                ("VALIDATION_DIMENSION_TOLERANCE_PCT", "5.0", "주요 치수 허용오차", "도메인/공정별 별도 profile 권장"),
                ("ENABLE_PARALLEL_PREVIEW", "true", "복수 preview 병렬", "VRAM/동시성 부족 시 false"),
            ],
            [2840, 1300, 2780, 2440],
            font_size=8.4,
        )

        self.heading("5. DesignState·GeometryContract·검증 등급", 1)
        self.heading("5.1 데이터 계약", 2)
        self.table(
            ["계약", "소유 정보", "왜 필요한가"],
            [
                ("DesignSpec", "category, units, components, features, relationships, dimensions", "VLM/규칙 결과를 실행 가능한 허용 스키마로 제한"),
                ("DesignState", "선택 2D, 목적, revision, 안전, 미확정, 일관성 우선순위", "모드 간 설계 의도와 변경 이력 유지"),
                ("GeometryContract", "primitive, center/size, material, requirement_id, hard constraints", "OpenSCAD/Blender/OpenUSD가 같은 구조를 생성"),
                ("ValidationReport", "치수·구성·관계·파일·grade", "자동 검증과 사람 승인을 분리"),
            ],
            [2200, 4120, 3040],
            font_size=8.4,
        )
        self.heading("5.2 일관성 우선순위", 2)
        self.numbers(
            [
                "기능과 동작 원리",
                "필수 구성요소와 수량",
                "주요 치수와 배치",
                "전체 비례와 실루엣",
                "재질·곡면·미세 외관",
            ]
        )
        self.heading("5.3 검증 등급", 2)
        self.table(
            ["코드", "UI 표시", "자동/사람", "허용 활용"],
            [
                ("concept", "컨셉 검토 가능", "자동", "외관 방향·아이디어"),
                ("structured", "구조 검토 가능", "자동", "배치·초기 엔지니어링·후속 CAD 입력"),
                ("validated", "자동 검증 완료", "자동", "계약을 통과한 pre-CAD 입력"),
                ("engineer_reviewed", "엔지니어 검토 완료", "인증 검토자", "상세 CAD 보완 진입"),
                ("manufacturing_approved", "제조 승인 완료", "승인 담당자", "공정·안전·공차 확인 후 제조"),
            ],
            [1900, 2200, 1600, 3660],
            font_size=8.5,
        )
        self.callout(
            "95% 해석",
            "SELF_FEEDBACK_TARGET=0.95는 내부 acceptance 목표다. 적어도 범주별 독립 홀드아웃 200건, "
            "Wilson 95% 하한, 엔지니어 평가가 필요하다. 자동 점수 하나로 제조 수율 95%를 주장하지 않는다.",
            tone="warning",
        )

    def _installation_and_operation(self) -> None:
        self.page_break()
        self.heading("6. 로컬 설치와 운영", 1)
        self.heading("6.1 선행 조건", 2)
        self.bullets(
            [
                "Windows 11 + Docker Desktop 또는 Ubuntu 22.04/24.04 + Docker Engine/Compose v2",
                "로컬 GPU 공급자를 쓸 경우 NVIDIA Driver, CUDA 호환 런타임, 충분한 VRAM",
                "ComfyUI는 기본 http://host.docker.internal:8188, TripoSR는 8081",
                "OpenSCAD와 Blender는 Worker 이미지에 설치하거나 호스트 native bridge를 사용",
                ".env는 .env.example에서 만들고 실제 비밀번호·Key를 Git에 넣지 않음",
            ]
        )
        self.heading("6.2 최초 실행", 2)
        self.code(
            "cd \"전체 풀스택\"\n"
            "Copy-Item .env.example .env\n"
            "# .env의 비밀번호와 로컬 provider URL을 설정\n"
            "docker compose -p xconcep up -d --build\n"
            "docker compose -p xconcep ps"
        )
        self.paragraph("기본 UI는 WEB_PORT 값이며 .env.example 기준 http://127.0.0.1:8080 이다.")
        self.heading("6.3 현재 검증용 포트 예시", 2)
        self.table(
            ["구성", "주소", "노출 원칙"],
            [
                ("Web", "http://127.0.0.1:18080", "검증 프로젝트 xconcep-e2e-r2에서 사용"),
                ("MySQL", "127.0.0.1:13307→3306", "검증용 loopback only"),
                ("ComfyUI", "host.docker.internal:8188", "사내/로컬만"),
                ("TripoSR", "host.docker.internal:8081", "사내/로컬만"),
                ("VLM verifier", "GPU-SERVER:8191", "Worker IP allowlist"),
            ],
            [2100, 3300, 3960],
            font_size=8.6,
        )
        self.heading("6.4 Health 확인", 2)
        self.code(
            "Invoke-RestMethod http://127.0.0.1:18080/api/system-status\n"
            "Invoke-RestMethod http://127.0.0.1:18080/health\n"
            "docker compose -p xconcep ps\n"
            "docker compose -p xconcep logs --tail 200 control-plane python-worker celery-worker"
        )
        self.heading("6.5 일상 운영 순서", 2)
        self.numbers(
            [
                "MySQL·Redis·Qdrant와 외부 공급자 health를 먼저 확인한다.",
                "ComfyUI/TripoSR 등 GPU 서비스를 시작하고 모델 이름을 상태 API에서 확인한다.",
                "Compose 애플리케이션을 시작한 뒤 내부 계정으로 로그인한다.",
                "짧은 smoke-live 또는 smoke-meeting-live를 1건 실행한다.",
                "queue 적체·GPU OOM·storage 사용량·품질 보고서를 확인한다.",
                "종료 전 실행 중 Job을 drain하고 DB·storage·Qdrant 백업 상태를 확인한다.",
            ]
        )
        self.heading("6.6 중지와 재기동", 2)
        self.code(
            "docker compose -p xconcep stop\n"
            "docker compose -p xconcep start\n"
            "# 컨테이너를 재생성하되 데이터 volume은 유지\n"
            "docker compose -p xconcep up -d --force-recreate"
        )
        self.callout(
            "데이터 주의",
            "docker compose down은 컨테이너를 내리지만 기본적으로 named volume을 지우지 않는다. "
            "down -v는 MySQL/Qdrant/Redis 데이터를 삭제하므로 복구 승인 없이 사용하지 않는다.",
            tone="danger",
        )

    def _configuration(self) -> None:
        self.page_break()
        self.heading("7. 주요 설정값", 1)
        self.paragraph(
            "표의 값은 .env.example과 docker-compose.yml의 현재 기본값이다. 운영값은 환경·GPU·데이터 권한에 "
            "맞춰 별도 Secret/Config로 관리한다."
        )
        self.heading("7.1 Web·DB·인증", 2)
        self.table(
            ["변수", "기본값", "역할", "운영 권장"],
            [
                ("WEB_PORT", "8080", "외부 UI port", "reverse proxy/TLS 뒤 배치"),
                ("MYSQL_HOST_PORT", "3307", "호스트 loopback port", "외부 공개 금지"),
                ("MYSQL_CONN_MAX_AGE", "60", "DB connection 재사용", "DB 부하 측정 후 조정"),
                ("AUTH_MODE", "internal_db", "disabled/internal_db/corporate_db", "현재 internal_db 유지"),
                ("AUTH_TOKEN_TTL_SECONDS", "28800", "세션 token 8시간", "정책에 맞춰 단축"),
                ("INTERNAL_AUTH_BOOTSTRAP_ENABLED", "true(example)", "테스트 계정 생성", "운영 전 false 검토"),
                ("CORS_ALLOW_ALL", "false", "CORS 전면 허용", "false 유지"),
                ("SYNC_PIPELINE", "false", "동기/비동기 실행", "운영 false + Celery"),
            ],
            [3000, 1900, 2520, 1940],
            font_size=8.0,
        )
        self.heading("7.2 외부 사내 MySQL 전환용 매핑", 2)
        self.table(
            ["변수 그룹", "입력 내용", "현재 처리"],
            [
                ("AUTH_DB_HOST/PORT/NAME/USER/PASSWORD", "외부 MySQL 연결·권한", "internal_db에서는 미사용"),
                ("AUTH_DB_TABLE", "직원 사용자 table, 기본 employees", "식별자 allowlist 검증"),
                ("*_COLUMN", "id/username/password/display/email/active", "실제 schema에 맞춰 매핑"),
                ("AUTH_DB_PASSWORD_SCHEME", "django_hash 등 검증 방식", "평문 금지"),
                ("AUTH_DB_SSL_CA", "TLS CA 파일", "원격 DB는 TLS 권장"),
                ("AUTH_DB_USERNAME_CASE_SENSITIVE", "아이디 대소문자 정책", "기본 false"),
            ],
            [3100, 3520, 2740],
            font_size=8.4,
        )
        self.callout(
            "현재 결정",
            "외부 사내 MySQL 인증은 추후에 연결한다. 지금은 내부 Docker MySQL과 bootstrap 계정으로만 테스트한다.",
            tone="info",
        )
        self.heading("7.3 2D 이미지", 2)
        self.table(
            ["변수", "기본값", "설명"],
            [
                ("OPENAI_IMAGE_MODE", "comfyui", "기본 local / 선택 openai / test mock"),
                ("OPENAI_IMAGE_MODEL", "gpt-image-2", "OpenAI 선택 모드 모델"),
                ("OPENAI_IMAGE_SIZE", "1536x1024", "OpenAI 요청 크기"),
                ("OPENAI_IMAGE_QUALITY", "medium", "OpenAI 품질"),
                ("OPENAI_IMAGE_MAX_REQUESTS_PER_DAY", "20", "일일 호출 상한"),
                ("COMFYUI_UNET_MODEL", "flux-2-klein-base-4b-fp8.safetensors", "로컬 FLUX UNET"),
                ("COMFYUI_CLIP_MODEL", "qwen_3_4b.safetensors", "텍스트 인코더"),
                ("COMFYUI_VAE_MODEL", "flux2-vae.safetensors", "VAE"),
                ("COMFYUI_WIDTH/HEIGHT", "1024/1024", "기본 1:1 concept"),
                ("COMFYUI_STEPS / CFG", "20 / 5.0", "sampling 조정"),
                ("IMAGE_CONCEPT_COUNT", "4", "한 요청 concept 수"),
                ("IMAGE_MIN_WIDTH/HEIGHT", "768/768", "기본 파일 gate"),
                ("IMAGE_MIN_FILE_BYTES", "10000", "빈/손상 파일 차단"),
                ("IMAGE_MIN_CHANNEL_STDDEV", "3.0", "blank 이미지 차단"),
            ],
            [3500, 2650, 3210],
            font_size=7.8,
        )
        self.heading("7.4 3D·품질", 2)
        self.table(
            ["변수", "기본값", "설명"],
            [
                ("SHAPE_PROVIDER", "triposr", "빠른 mesh 공급자"),
                ("SHAPE_TIMEOUT_SECONDS", "1800", "이미지→3D 제한"),
                ("OPENSCAD_MODE / BIN", "auto / openscad", "native 우선, fallback 가능"),
                ("OPENSCAD_TIMEOUT_SECONDS", "600", "구조 생성 제한"),
                ("BLENDER_MODE / BIN", "auto / blender", "native bridge"),
                ("BLENDER_TIMEOUT_SECONDS", "1800", "render/export 제한"),
                ("SELF_FEEDBACK_TARGET", "0.95", "내부 acceptance 목표"),
                ("SELF_FEEDBACK_MAX_ATTEMPTS", "3", "부분 재생성 최대 반복"),
                ("VALIDATION_DIMENSION_TOLERANCE_PCT", "5.0", "기본 치수 허용오차"),
            ],
            [3660, 2300, 3400],
            font_size=8.2,
        )
        self.heading("7.5 회의 음성·OpenUSD", 2)
        self.table(
            ["변수", "기본값", "설명"],
            [
                ("SPEECH_MODE", "mock(example)", "실제 검증은 faster_whisper"),
                ("WHISPER_MODEL", "large-v3-turbo(example)", "검증 환경은 small"),
                ("WHISPER_DEVICE/COMPUTE_TYPE", "auto/auto", "CUDA 우선, 실패 시 CPU 재시도"),
                ("DIARIZATION_MODE", "none", "화자 분리 비활성"),
                ("MAX_AUDIO_UPLOAD_MB", "30", "chunk upload 제한"),
                ("OPENUSD_GENERATE_USDC", "true", "binary stage 생성"),
                ("OMNIVERSE_GENERATE_LAYERS", "true", "root/geometry/looks/meeting"),
                ("OMNIVERSE_ENABLE_PHYSICS", "true", "physics metadata"),
                ("OMNIVERSE_NUCLEUS_URL", "empty", "운영 Nucleus 미연결"),
                ("OMNIVERSE_STREAM_URL", "empty", "WebRTC 미연결"),
            ],
            [3460, 2480, 3420],
            font_size=8.2,
        )

    def _server_migration(self) -> None:
        self.page_break()
        self.heading("8. 전체 서버 이식", 1)
        self.heading("8.1 권장 대상 구조", 2)
        self.table(
            ["계층", "권장 위치", "비고"],
            [
                ("Web/Control/Data", "사내 VM 또는 Kubernetes node", "MySQL/Qdrant/Redis private network"),
                ("ComfyUI/TripoSR/Whisper", "로컬 또는 GPU inference server", "공급자별 queue·VRAM 분리"),
                ("CAD VLM 학습", "별도 고성능 GPU server", "학습 중 verifier 동시 실행 금지"),
                ("VLM verifier", "승인 adapter를 가진 inference server", "8191을 Worker IP에만 허용"),
                ("Artifact storage", "NAS/object storage 또는 공유 volume", "파일 잠금·버전·백업 필요"),
            ],
            [2300, 3520, 3540],
            font_size=8.5,
        )
        self.heading("8.2 이식 전 freeze", 2)
        self.numbers(
            [
                "Git commit/working-tree diff와 현재 .env 변수 목록을 기록한다. 비밀값은 별도 Secret로 이동한다.",
                "docker compose config 결과와 사용 이미지 digest를 저장한다.",
                "MySQL 논리 백업·Qdrant snapshot·storage 파일 스냅샷을 같은 release ID로 묶는다.",
                "ComfyUI workflow/model, TripoSR model, Whisper model, CAD VLM base revision/adapter SHA를 기록한다.",
                "scripts/test-local.ps1과 live smoke의 기준 보고서·스크린샷을 보존한다.",
            ]
        )
        self.code(
            "python scripts/capture-quality-environment.py\n"
            "python scripts/verify-mysql-backup.py\n"
            "docker compose config > compose.resolved.yml\n"
            "python training/cad-vlm/scripts/export_bundle.py"
        )
        self.heading("8.3 애플리케이션 복사", 2)
        self.bullets(
            [
                "저장소 소스는 Git bundle/private remote 또는 승인된 ZIP으로 복사한다.",
                ".env, storage, outputs, 모델 캐시는 소스 ZIP과 분리한다.",
                "대상에서 파일 SHA-256과 portable bundle의 bundle-manifest.json을 모두 검증한다.",
                "Windows 개발 폴더를 그대로 Linux에 복사할 때 CRLF, 실행 권한, 대소문자 파일명을 확인한다.",
            ]
        )
        self.heading("8.4 대상 서버 설치", 2)
        self.code(
            "git clone <internal-repository> xconcep-ai\n"
            "cd xconcep-ai/\"전체 풀스택\"\n"
            "cp .env.example .env\n"
            "# Secret와 provider URL 설정\n"
            "docker compose pull\n"
            "docker compose build\n"
            "docker compose up -d"
        )
        self.heading("8.5 데이터 복구 순서", 2)
        self.numbers(
            [
                "새 MySQL을 빈 상태로 기동하고 schema migration 버전을 맞춘다.",
                "검증한 논리 dump를 복원하고 테이블별 row count를 원본 보고서와 비교한다.",
                "storage를 원래 상대 경로 구조로 복원하고 표본 GLB/PNG/USD를 해시·재개방한다.",
                "Qdrant snapshot을 복원하고 collection count와 search smoke를 확인한다.",
                "Redis는 실행 중 Job을 이식하지 말고 queue drain 후 새로 시작한다.",
            ]
        )
        self.heading("8.6 공급자 연결", 2)
        self.table(
            ["공급자", "대상 설정", "연결 확인"],
            [
                ("ComfyUI", "COMFYUI_BASE_URL, model 파일명", "health + 1 concept"),
                ("TripoSR", "SHAPE_API_URL", "/health + GLB magic/size"),
                ("Whisper/NIM", "SPEECH_MODE, model/URL", "한국어 fixture recall·치수"),
                ("vLLM", "VLLM_BASE_URL, model name", "/v1/models + 요구 구조화"),
                ("CAD VLM verifier", "IMAGE_SEMANTIC_VERIFIER_URL/API_KEY", "/health + /verify"),
                ("Nucleus/WebRTC", "OMNIVERSE_* URL", "운영 환경에서 별도 검증"),
            ],
            [2100, 3850, 3410],
            font_size=8.4,
        )
        self.heading("8.7 이식 승인 Gate", 2)
        self.numbers(
            [
                "모든 service health가 green이고 내부 인증 login이 가능함.",
                "로컬 회귀 100건과 CAD 패키지 15건이 모두 통과함.",
                "직접 입력 live E2E와 한국어 회의 live E2E가 통과함.",
                "UI에서 최신 Project, 2D, 3D, 다운로드 링크, viewer camera가 정상임.",
                "원본과 대상의 데이터 row count·artifact SHA·모델 revision을 비교함.",
                "기능 Gate와 0.95 품질 Gate를 별도 기록하고 후자가 실패하면 제조 승격하지 않음.",
            ]
        )
        self.heading("8.8 롤백", 2)
        self.code(
            "# 애플리케이션: 직전 승인 image/tag와 .env snapshot으로 재기동\n"
            "docker compose up -d --force-recreate\n"
            "# CAD VLM: VLM_MODEL을 직전 승인 adapter로 변경\n"
            "docker compose --profile serve up -d --force-recreate verifier\n"
            "# 긴급 verifier 우회: IMAGE_SEMANTIC_VERIFIER_URL을 비우고 Worker 재기동"
        )

    def _libraries_and_paths(self) -> None:
        self.page_break()
        self.heading("9. 사용 라이브러리와 저장소 구조", 1)
        self.heading("9.1 직접 의존성", 2)
        self.table(
            ["영역", "핵심 라이브러리", "현재 버전/범위", "역할"],
            [
                ("Control", "Django / DRF", "5.2.16 / 3.17.1", "API·ORM·인증"),
                ("Control", "Celery / Redis client", "5.5.3 / 6.2.0", "비동기 Job"),
                ("Control", "mysqlclient", "2.2.7", "MySQL 연결"),
                ("Worker", "FastAPI / Uvicorn", "0.116.1 / 0.35.0", "AI 내부 API"),
                ("Worker", "Pillow / NumPy / SciPy", "11.3 / 2.3.1 / ≥1.14", "이미지·수치 처리"),
                ("Worker", "trimesh", "4.7.1", "mesh 재개방·검증"),
                ("Worker", "usd-core", "≥24.11", "USDA/USDC 작성·재개방"),
                ("Speech", "faster-whisper", "≥1.1.1", "로컬 STT 선택"),
                ("Knowledge", "qdrant-client", "1.18.0", "RAG vector adapter"),
                ("Knowledge", "pypdf", "5.9.0", "PDF fallback 추출"),
                ("Agent", "NVIDIA NAT", "≥1.8,<2 선택", "tool workflow profile"),
                ("3D service", "TripoSR stack", "transformers 4.46 등", "single-view mesh"),
                ("Frontend", "PHP", "container runtime", "UI·reverse proxy"),
                ("Omniverse Web", "@nvidia/ov-web-rtc", "6.4.4", "WebRTC client 선택"),
                ("CAD VLM", "transformers / TRL", "4.57.3 / 0.22.2", "Qwen3-VL 학습"),
                ("CAD VLM", "Unsloth / PEFT", "2026.7.5 / ≥0.14", "LoRA/QLoRA"),
                ("CAD VLM", "xformers", "0.0.31.post1", "attention 최적화"),
                ("Tracking", "TensorBoard / Trackio", "≥2.18 / ≥0.3", "학습 지표"),
                ("Audio fixture", "edge-tts / piper-onnx", "7.2.8 / 1.0.6", "QA 음성 생성 전용"),
            ],
            [1600, 2680, 2080, 3000],
            font_size=7.5,
        )
        self.heading("9.2 Native/infra", 2)
        self.table(
            ["구성", "버전/모드", "역할", "주의"],
            [
                ("MySQL", "8.4", "상태·내부 인증", "UTF8MB4, 외부 공개 금지"),
                ("Qdrant", "1.18.2", "RAG memory", "snapshot 필요"),
                ("Redis", "8.2-alpine", "Celery queue", "AOF enabled"),
                ("OpenSCAD", "native binary", "parametric SCAD→STL/GLB", "strict native acceptance 권장"),
                ("Blender", "native binary", "assembly/render/OpenUSD", "버전별 exporter 차이"),
                ("Docker Compose", "v2", "서비스 재현", "프로젝트명 고정"),
                ("CUDA/PyTorch", "서버별 호환", "학습·추론", "driver/SM 지원 확인"),
            ],
            [2000, 2100, 2800, 2460],
            font_size=8.3,
        )
        self.heading("9.3 주요 디렉터리", 2)
        self.table(
            ["경로", "내용"],
            [
                ("frontend-php/", "웹 UI, 인증 세션, API proxy, Three.js viewer"),
                ("control-plane-drf/", "Project/Job/Meeting/Auth/QualityEvidence"),
                ("agent-layer-nat/", "결정론 gateway와 NAT plugin/profile"),
                ("python-worker/", "분석·라우팅·2D·3D·품질·OpenUSD"),
                ("knowledge-service/", "Qdrant ingest/search와 선택 NeMo Retriever"),
                ("triposr-service/", "로컬 image-to-3D service"),
                ("omniverse-rtx/, omniverse-web-client/", "RTX/WebRTC 선택 배포"),
                ("training/cad-vlm/", "PHP CAD importer·통합 데이터·Qwen3-VL 학습·verifier"),
                ("scripts/", "회귀·live smoke·fixture·품질·백업 자동화"),
                ("storage/projects/<ID>/", "Project별 입력·분석·2D·3D·USD·보고서"),
                ("docs/", "설계·운영·검증 문서와 이 Word 매뉴얼"),
            ],
            [3300, 6060],
            font_size=8.6,
        )

    def _training(self, training: Path) -> None:
        self.page_break()
        self.heading("10. CAD VLM 학습 방법", 1)
        self.picture(training, "그림 4. PHP DXF/STEP 전처리부터 운영 승격까지", width=6.4)
        self.callout(
            "학습 책임 경계",
            "Qwen3-VL은 다중 시점 이미지와 요구문에서 DesignSpec JSON을 예측한다. "
            "STEP/OpenSCAD/Python 코드를 직접 생성·실행하지 않으며, JSON 검증 후 기존 결정론 생성기로 넘긴다.",
            tone="info",
        )
        self.heading("10.1 서버 권장 사양", 2)
        self.table(
            ["목적", "GPU", "RAM", "디스크", "설정"],
            [
                ("코드·데이터 dry-run", "불필요", "16GB", "10GB", "--dry-run"),
                ("4B QLoRA 개발", "VRAM ≥24GB", "64GB", "150GB", "qwen3-vl-4b-qlora.json"),
                ("8B QLoRA", "VRAM ≥48GB", "128GB", "250GB", "qwen3-vl-8b-qlora.json"),
                ("8B BF16 LoRA", "H100/A100 80GB", "≥128GB", "300GB", "qwen3-vl-8b-h100-lora.json"),
                ("2–4 GPU", "동일 GPU 권장", "256GB", "≥500GB", "accelerate launch"),
            ],
            [1900, 1900, 1200, 1200, 3160],
            font_size=8.2,
        )
        self.heading("10.2 PHP CAD 전처리 패키지 입력", 2)
        self.paragraph(
            "첨부된 CAD AI Dataset Studio PHP 프로그램의 패키지 구조를 입력으로 사용한다. "
            "각 sample은 manifest.json, geometry/geometry.json, quality/report.json, metadata/label.json을 가져야 한다."
        )
        self.table(
            ["입력 파일", "필수 정보", "Importer 처리"],
            [
                ("manifest.json", "sample_id, source format/hash, split, provenance", "ID·split·source 기록"),
                ("geometry/geometry.json", "bbox, entity_counts, topology, surfaces, primitive/points", "관측 CAD context + PNG preview"),
                ("quality/report.json", "score", "minimum-quality 미만 제외"),
                ("metadata/label.json", "category, training_prompt/description", "part/module/equipment 매핑"),
            ],
            [2450, 3480, 3430],
            font_size=8.5,
        )
        self.heading("10.3 PHP 패키지 import", 2)
        self.code(
            "cd training/cad-vlm\n"
            "python scripts/import_php_cad_dataset.py \\\n"
            "  --input data/incoming/php-packages.zip \\\n"
            "  --output data/raw/php-cad-v1 \\\n"
            "  --license LicenseRef-Company-Approved \\\n"
            "  --training-allowed \\\n"
            "  --minimum-quality 0.90 \\\n"
            "  --category-map '{\"bracket\":\"part\",\"cell\":\"equipment\"}'"
        )
        self.bullets(
            [
                "--training-allowed를 명시하지 않으면 import 자체가 실패한다.",
                "ZIP의 절대 경로와 .. traversal member를 차단한다.",
                "중복 sample_id, 유효하지 않은 JSON, 품질 미달, 허용되지 않은 category map을 차단한다.",
                "manifest와 geometry SHA-256을 cad_context/provenance에 남긴다.",
            ]
        )
        self.heading("10.4 기존 데이터 방식도 통합 스키마로 변환", 2)
        self.code(
            "python scripts/preprocess_dataset.py \\\n"
            "  --input data/raw/php-cad-v1/records.jsonl \\\n"
            "  --output data/production-v1 \\\n"
            "  --max-image-side 2048 \\\n"
            "  --min-images 1 \\\n"
            "  --split 0.8,0.1,0.1\n"
            "python scripts/validate_dataset.py --dataset data/production-v1"
        )
        self.paragraph(
            "기존 원본 JSONL도 같은 preprocess_dataset.py를 사용한다. 동일 프롬프트, 동일 이미지, 같은 CAD에서 "
            "파생된 뷰는 같은 split에 묶어 train/eval/test 누수를 막는다."
        )
        self.heading("10.5 통합 records.jsonl", 2)
        self.code(
            '{\n'
            '  "schema_version": "xconcep.cad-vlm-sample/1.0",\n'
            '  "id": "equipment_conveyor_inspection_001",\n'
            '  "category": "equipment", "split": "train",\n'
            '  "prompt": "폭 1600mm ...",\n'
            '  "images": [{"path":"images/front.png","view":"front"}],\n'
            '  "design_spec": {}, "geometry_contract": {},\n'
            '  "cad_context": {"schema":"xconcep.php-cad-context/1.0"},\n'
            '  "provenance": {"license":"LicenseRef-Company-Approved",\n'
            '                 "training_allowed":true}\n'
            '}'
        )
        self.heading("10.6 샘플 데이터 재생성", 2)
        self.code(
            "python scripts/build_sample_dataset.py --output data/examples --replace\n"
            "python scripts/validate_dataset.py --dataset data/examples"
        )
        self.paragraph(
            "현재 샘플은 PHP 호환 패키지 흐름으로 재생성된 9 records·9 observed PNG이며 train 6/eval 3이다. "
            "배선 확인용일 뿐 성능 학습용이 아니다. 생산 데이터는 범주별 독립 설계 수천 건이 필요하다."
        )
        self.heading("10.7 설치와 dry-run", 2)
        self.code(
            "# Ubuntu 권장\n"
            "chmod +x scripts/*.sh\n"
            "PYTORCH_INDEX_URL=https://download.pytorch.org/whl/cu128 ./scripts/install-server.sh\n"
            "source .venv/bin/activate\n"
            "python scripts/check_server.py\n"
            "python scripts/train_vlm.py --config configs/qwen3-vl-4b-qlora.json --dry-run"
        )
        self.heading("10.8 학습 실행", 2)
        self.code(
            "./scripts/train.sh configs/qwen3-vl-4b-qlora.json\n"
            "# 다중 GPU 예시\n"
            "accelerate launch --num_processes 4 scripts/train_vlm.py \\\n"
            "  --config configs/qwen3-vl-8b-h100-lora.json"
        )
        self.heading("10.9 주요 조정값", 2)
        self.table(
            ["값", "4B QLoRA", "8B H100 LoRA", "조정 효과/주의"],
            [
                ("load_in_4bit", "true", "false", "VRAM 절감 vs BF16 정밀도"),
                ("batch × accumulation", "1 × 32", "2 × 16", "유효 batch 32, GPU 수 포함해 기록"),
                ("learning_rate", "8e-5", "4e-5", "과적합/불안정 시 낮춤"),
                ("epochs", "2", "2", "독립 eval 악화 시 중단"),
                ("LoRA r/alpha", "32/64", "64/128", "용량 증가, VRAM/과적합 증가"),
                ("optimizer", "adamw_8bit", "adamw_torch_fused", "환경 호환 확인"),
                ("warmup_ratio", "0.03", "0.03", "초기 발산 방지"),
                ("weight_decay", "0.01", "0.01", "과적합 완화"),
                ("max_grad_norm", "1.0", "1.0", "gradient 폭주 완화"),
                ("max_pixels", "null", "null", "OOM 시 먼저 제한"),
                ("max_images", "3", "3", "정면/상면/우측면 기본"),
                ("seed", "3407", "3407", "재현 비교용, bit-exact 보장 아님"),
                ("resume", "auto", "auto", "최신 checkpoint 자동 재개"),
            ],
            [2450, 1650, 1900, 3360],
            font_size=7.8,
        )
        self.heading("10.10 실제 1-step 호환성 검증", 2)
        self.table(
            ["항목", "결과"],
            [
                ("설정", "configs/qwen3-vl-4b-php-cad-smoke.json"),
                ("데이터", "data/php-api-e2e-smoke-v1 · train 1 record"),
                ("실행", "Qwen3-VL 4B · LoRA r=8 · max_steps=1"),
                ("결과", "train_loss=3.2973738 · runtime=39.1659s"),
                ("해석", "학습 코드·데이터 collator·checkpoint가 실행됨. 모델 품질 증거는 아님."),
            ],
            [2600, 6760],
            font_size=8.7,
        )
        self.heading("10.11 모니터링", 2)
        self.code(
            "tensorboard --logdir outputs --host 0.0.0.0 --port 6006\n"
            "# 외부 전송 금지 환경\n"
            '# config tracking.report_to를 ["tensorboard"]로 설정'
        )
        self.bullets(
            [
                "train/eval loss, learning rate, gradient norm, tokens/sec, GPU 사용률, OOM을 함께 본다.",
                "모델 크기·데이터 version·seed·GPU 수·유효 batch·commit hash를 run metadata에 기록한다.",
                "한 실험에서 한 변수만 바꾸고 run_name을 다르게 한다.",
                "loss만으로 구성요소·치수·제조 적중률을 판단하지 않는다.",
            ]
        )
        self.heading("10.12 추론과 운영 연결", 2)
        self.code(
            "python scripts/infer_vlm.py \\\n"
            "  --model outputs/qwen3-vl-4b-designspec/adapter \\\n"
            "  --image front.png --image top.png --image right.png \\\n"
            "  --prompt \"폭 1600mm ...\" --output prediction.json\n"
            "docker compose --profile serve up -d --build verifier\n"
            "curl http://127.0.0.1:8191/health"
        )
        self.heading("10.13 외부 학습 서버 이식", 2)
        self.code(
            "python scripts/export_bundle.py\n"
            "sha256sum -c xconcep-cad-vlm-portable.zip.sha256\n"
            "unzip xconcep-cad-vlm-portable.zip\n"
            "cd xconcep-cad-vlm\n"
            "python3 scripts/verify_install.py\n"
            "./scripts/install-server.sh"
        )
        self.paragraph(
            "portable ZIP에는 모델 weight, .env, outputs, HF cache를 포함하지 않는다. Adapter만 운영 서버로 "
            "이동할 때는 adapter_config.json, safetensors, processor/tokenizer, base model revision을 한 release로 묶는다."
        )
        self.heading("10.14 95% 승격 조건", 2)
        self.code(
            "python scripts/evaluate_predictions.py \\\n"
            "  --dataset data/production-v1 \\\n"
            "  --predictions outputs/predictions-v1 \\\n"
            "  --split test --dimension-tolerance-pct 5 \\\n"
            "  --target 0.95 --min-cases-per-category 200 \\\n"
            "  --output outputs/holdout-report-v1.json"
        )
        self.table(
            ["검사", "통과 조건"],
            [
                ("JSON/스키마", "object, category 일치, units=mm, 허용 key/kind"),
                ("구성요소", "필수 kind recall 100%, 요구 수량 일치"),
                ("feature", "필수 hole/rib/slot recall 100%, 수량 일치"),
                ("치수", "정답의 모든 주요 치수가 허용오차 이내"),
                ("관계", "필수 subject-relation-object recall 100%"),
                ("생성 E2E", "GeometryContract→native CAD→GLB/STL/USD 재개방"),
                ("통계", "part/module/equipment 각 ≥200, Wilson 95% 하한 ≥0.95"),
                ("사람 검토", "고정 blind holdout에서 설비 엔지니어 승인"),
            ],
            [2600, 6760],
            font_size=8.6,
        )

    def _qa_e2e(self) -> None:
        self.page_break()
        self.heading("11. QA와 E2E 테스트 방법", 1)
        self.heading("11.1 테스트 계층", 2)
        self.table(
            ["계층", "목적", "외부 연결", "실패 시 조치"],
            [
                ("정적/구문", "PHP·JS·Python 기본 오류", "없음", "즉시 수정"),
                ("Unit", "parser·router·contract·security", "mock", "해당 module 수정"),
                ("Service", "DRF/Worker/Agent/Knowledge API", "local", "API·schema 확인"),
                ("Native CAD", "OpenSCAD/Blender/OpenUSD 실제 파일", "native binary", "환경/timeout/format 확인"),
                ("Live provider", "ComfyUI·TripoSR·Whisper", "local GPU", "model/URL/VRAM 확인"),
                ("Browser UI", "login·입력·이력·viewer·다운로드", "전체 stack", "DOM/console/network 확인"),
                ("Training smoke", "dataset→collator→LoRA→checkpoint", "GPU/HF model", "버전·CUDA·EOS·cache 확인"),
                ("품질/통계", "외관·치수·구성 독립 holdout", "평가기/사람", "승격 보류"),
            ],
            [1750, 2860, 1900, 2850],
            font_size=8.1,
        )
        self.heading("11.2 전체 로컬 회귀", 2)
        self.code(
            "cd \"전체 풀스택\"\n"
            "powershell -ExecutionPolicy Bypass -File .\\scripts\\test-local.ps1 -SkipCompose"
        )
        self.table(
            ["Suite", "현재 결과"],
            [
                ("control-plane-drf", "12 passed"),
                ("python-worker", "85 passed"),
                ("agent-layer", "1 passed"),
                ("knowledge-service", "2 passed"),
                ("frontend", "PHP/JS syntax PASS"),
                ("합계", "100 tests + syntax"),
            ],
            [4200, 5160],
            font_size=8.8,
        )
        self.heading("11.3 CAD 학습 패키지 회귀", 2)
        self.code(
            "cd training/cad-vlm\n"
            "python -m pytest -q\n"
            "# 현재: 15 passed"
        )
        self.callout(
            "Windows 실행기 주의",
            "이번 검증에서 Windows py.exe는 ‘logon session does not exist’로 실패했지만 동일 코드가 "
            "저장소 .venv Python에서 15/15 통과했다. CI와 매뉴얼 명령은 활성 venv의 python을 명시한다.",
            tone="warning",
        )
        self.heading("11.4 직접 입력 Live E2E", 2)
        self.code(
            "powershell -ExecutionPolicy Bypass -File .\\scripts\\smoke-live.ps1 \\\n"
            "  -BaseUrl http://127.0.0.1:18080\n"
            "# 외관 목표까지 강제할 때만 추가\n"
            "# -RequireQualityTarget"
        )
        self.paragraph(
            "smoke-live.ps1은 내부 DB로 로그인하고 Project/Job을 polling하며 2D 4안, 선택 3D, GLB 크기, "
            "검증 등급, 품질 점수를 JSON으로 출력한다. -RequireQualityTarget 없이도 기능 E2E는 엄격히 검사한다."
        )
        self.heading("11.5 한국어 회의 음성 Fixture 생성", 2)
        self.code(
            "# 온라인 고품질 QA 음성 (최종 사용)\n"
            "..\\.e2e-tts-venv\\Scripts\\python.exe \\\n"
            "  scripts\\generate-korean-edge-meeting-fixtures.py\n"
            "# 오프라인 Piper 대안\n"
            "..\\.e2e-tts-venv\\Scripts\\python.exe \\\n"
            "  scripts\\generate-korean-meeting-fixtures.py"
        )
        self.paragraph(
            "최종 fixture는 Microsoft 한국어 SunHi/InJoon 음성을 edge-tts 7.2.8로 생성한다. "
            "재배포 전 Microsoft 서비스 약관을 검토한다. Piper KSS fixture는 오프라인 대안이지만 이번 "
            "Faster-Whisper에서 핵심어 recall 0.0833으로 부적합하여 최종 기준에서 제외했다."
        )
        self.heading("11.6 회의 음성 Live E2E", 2)
        self.code(
            "powershell -ExecutionPolicy Bypass -File .\\scripts\\smoke-meeting-live.ps1 \\\n"
            "  -BaseUrl http://127.0.0.1:18080 \\\n"
            "  -FixtureDir .\\storage\\e2e-audio\\korean-industrial-meeting-edge"
        )
        self.table(
            ["검사", "Pass 조건", "2026-07-23 결과"],
            [
                ("인증/Project", "internal_db login, Meeting 생성", "PASS"),
                ("Audio", "3 chunk 업로드·전사", "PASS"),
                ("Keyword recall", "≥0.75", "0.9167"),
                ("Dimensions", "width/depth/height 전부 존재", "1200/800/1600"),
                ("2D", "concepts=4", "4"),
                ("3D route", "openscad_equipment", "PASS"),
                ("Validation", "grade=structured", "PASS"),
                ("GLB", "유효 header + 최소 크기", "54,884 bytes"),
                ("Quality target", "score≥0.95", "0.4161 · FAIL"),
            ],
            [2150, 3440, 3770],
            font_size=8.3,
        )
        self.heading("11.7 브라우저 UI QA", 2)
        self.numbers(
            [
                "내부 계정 login 후 사용자 표시명과 logout 버튼을 확인한다.",
                "직접 입력/회의 음성 tab, category, result mode, advanced route를 전환한다.",
                "Project 생성 후 2D 4안과 선택 상태를 확인한다.",
                "3D 결과에서 선택 2D, route, grade, 재생성 버튼을 확인한다.",
                "ISO→FRONT→TOP→ISO를 눌러 각 버튼의 pressed/active 상태와 camera 변화를 확인한다.",
                "GLB, STL, SCAD, Geometry JSON, USDA, USDC, package, manifest, PNG 링크를 확인한다.",
                "작업 이력을 열고 최신 Project가 3D 완료로 로드되는지 확인한다.",
                "스크린샷, console error, failed network 요청을 evidence 폴더에 기록한다.",
            ]
        )
        self.heading("11.8 산출물 정적 검증", 2)
        self.table(
            ["파일", "검사"],
            [
                ("PNG", "decode, width/height, bytes, channel stddev, aspect"),
                ("GLB", "magic=glTF, scene load, finite vertex/face, bbox"),
                ("STL", "mesh 재개방, positive extent, NaN 없음"),
                ("SCAD", "native OpenSCAD exit=0, STL 생성"),
                ("USDA/USDC", "pxr Stage.Open, default prim, up axis, metersPerUnit"),
                ("OpenUSD package", "sublayer/reference/payload/texture 경로, manifest"),
                ("Geometry JSON", "contract SHA, requirement coverage, finite dimensions"),
            ],
            [2600, 6760],
            font_size=8.6,
        )
        self.heading("11.9 품질 프로그램", 2)
        self.code(
            "python scripts/compare-quality-baseline.py --contract quality/advanced-baseline.json\n"
            "python scripts/analyze-evaluation-reliability.py \\\n"
            "  --minimum-holdout-cases 120 --minimum-seeds 3 --minimum-score-pct 95\n"
            "python scripts/benchmark-cad-roundtrip.py\n"
            "python scripts/benchmark-openusd-advanced.py"
        )
        self.bullets(
            [
                "calibration과 holdout을 고정하고 prompt/image SHA와 seed를 기록한다.",
                "평가기 선택과 채점에 같은 모델을 쓰면 independent_evaluation=false로 표시한다.",
                "평균 점수뿐 아니라 seed별 최저, 분산, Wilson/신뢰구간을 기록한다.",
                "사람 검수 또는 공식 평가가 없으면 대체 평가를 독립 평가라고 부르지 않는다.",
            ]
        )
        self.heading("11.10 테스트 증거 최소 항목", 2)
        self.table(
            ["분류", "필수 기록"],
            [
                ("코드", "commit/diff, suite, 명령, exit code, pass count"),
                ("환경", "OS, Docker, GPU/driver, Python/package, native tool"),
                ("모델", "provider, model/revision, workflow SHA, config"),
                ("입력", "prompt/audio/image SHA, category, expected contract"),
                ("출력", "Project ID, artifact path/size/SHA, validation report"),
                ("품질", "score, target, independent 여부, 실패 항목"),
                ("UI", "screenshot, 주요 DOM 상태, console/network 오류"),
            ],
            [2200, 7160],
            font_size=8.7,
        )

    def _troubleshooting(self) -> None:
        self.page_break()
        self.heading("12. 장애 대응", 1)
        self.table(
            ["증상", "가능 원인", "확인", "조치"],
            [
                ("401/로그인 실패", "AUTH_MODE/계정/DB", "/api/auth/config, DB health", "bootstrap·비밀번호·token 확인"),
                ("Job이 계속 pending", "Celery/Redis/Worker", "celery logs, Redis ping", "worker 재기동·queue 확인"),
                ("ComfyUI timeout", "모델명/VRAM/queue", "8188 health·workflow", "모델명·timeout·VRAM mode"),
                ("2D blank/손상", "workflow output", "manifest quality checks", "seed 재생성·모델/노드 수정"),
                ("TripoSR 연결 실패", "uvicorn/URL", "8081 /health", "서비스 재기동·host route"),
                ("3D가 blob", "single-view 한계", "self feedback·selected 2D", "전문 OpenSCAD/Blender로 route"),
                ("OpenSCAD timeout", "복잡도/native 미설치", "binary path·stderr", "INSTALL_OPENSCAD·scope 축소"),
                ("Blender export 실패", "버전 API 차이", "script report", "호환 fallback·timeout"),
                ("STT keyword 낮음", "음질/발음/model", "transcript·recall", "fixture/Whisper model/noise 개선"),
                ("한글 치수 누락", "한국어 수사 parser", "analysis dimensions", "천/백 수사 normalize 확인"),
                ("USDC 미생성", "usd-core/pxr", "worker health·logs", "usd-core 설치·flag"),
                ("학습 CUDA OOM", "pixels/batch/rank", "nvidia-smi·trace", "max_pixels→batch→rank 순으로 축소"),
                ("학습 EOS 경고", "VLM collator/tokenizer", "sample labels", "현재 train_vlm collator 유지"),
                ("py.exe 실행 실패", "Windows session/launcher", "python -V", "활성 venv python.exe 명시"),
                ("외관 0.95 실패", "평가기/형상 불일치", "self_feedback_report", "승격 보류·학습/생성기 개선"),
            ],
            [1900, 2300, 2360, 2800],
            font_size=7.7,
        )
        self.heading("12.1 로그 우선순위", 2)
        self.numbers(
            [
                "UI Network/Console에서 실패 endpoint와 status를 확인한다.",
                "control-plane에서 Project/Job 상태·예외를 확인한다.",
                "celery-worker에서 task retry/timeout을 확인한다.",
                "agent-layer에서 workflow 전달을 확인한다.",
                "python-worker에서 provider 응답·native stderr·validation을 확인한다.",
                "공급자(ComfyUI/TripoSR/vLLM/Whisper) 자체 로그를 확인한다.",
                "storage의 manifest/report를 최종 근거로 대조한다.",
            ]
        )
        self.heading("12.2 문제 보고 템플릿", 2)
        self.code(
            "Release/commit:\n"
            "Project ID / Job ID:\n"
            "실행 명령과 시각:\n"
            "입력 SHA / category / route:\n"
            "기대 결과 / 실제 결과:\n"
            "HTTP status / 로그 첫 원인:\n"
            "artifact·report·screenshot 경로:\n"
            "기능 Gate / 품질 Gate:"
        )

    def _cli_reference(self) -> None:
        self.heading("13. CLI 빠른 참조", 1)
        self.table(
            ["목적", "명령"],
            [
                ("전체 시작", "docker compose -p xconcep up -d --build"),
                ("상태", "docker compose -p xconcep ps"),
                ("로그", "docker compose -p xconcep logs --tail 200 <service>"),
                ("로컬 회귀", "powershell -File scripts/test-local.ps1 -SkipCompose"),
                ("직접 E2E", "powershell -File scripts/smoke-live.ps1 -BaseUrl http://127.0.0.1:18080"),
                ("음성 E2E", "powershell -File scripts/smoke-meeting-live.ps1 -BaseUrl ... -FixtureDir ..."),
                ("Edge 음성", "python scripts/generate-korean-edge-meeting-fixtures.py"),
                ("CAD tests", "cd training/cad-vlm; python -m pytest -q"),
                ("샘플 dataset", "python scripts/build_sample_dataset.py --output data/examples --replace"),
                ("PHP CAD import", "python scripts/import_php_cad_dataset.py --input ... --output ... --license ... --training-allowed"),
                ("dataset normalize", "python scripts/preprocess_dataset.py --input ... --output ..."),
                ("dataset validate", "python scripts/validate_dataset.py --dataset data/production-v1"),
                ("학습 dry-run", "python scripts/train_vlm.py --config configs/qwen3-vl-4b-qlora.json --dry-run"),
                ("학습", "./scripts/train.sh configs/qwen3-vl-4b-qlora.json"),
                ("홀드아웃", "python scripts/evaluate_predictions.py --dataset ... --predictions ... --split test"),
                ("portable ZIP", "python scripts/export_bundle.py"),
                ("verifier", "docker compose --profile serve up -d --build verifier"),
                ("환경 lock", "python scripts/capture-quality-environment.py"),
                ("품질 baseline", "python scripts/compare-quality-baseline.py --contract quality/advanced-baseline.json"),
                ("MySQL 복구 drill", "python scripts/verify-mysql-backup.py"),
            ],
            [2600, 6760],
            font_size=7.7,
        )

    def _design_rationale(self) -> None:
        self.page_break()
        self.heading("14. 주요 설계 이유와 트레이드오프", 1)
        self.table(
            ["결정", "선택 이유", "비용/한계", "후속 고도화"],
            [
                ("ComfyUI/FLUX 기본", "로컬·비밀 유지·반복 비용 통제", "GPU/모델 운영", "독립 의미 평가·workflow versioning"),
                ("OpenAI Image 추가", "기존 프롬프트 제한과 비교 가능한 고품질 option", "외부 전송·비용", "승인된 요청만·budget gate"),
                ("전문 OpenSCAD + 범용 유지", "치수·구조 재현성과 fallback 동시 확보", "자유 곡면 약함", "부품/모듈/설비 template 확장"),
                ("TripoSR 유지", "빠른 mesh preview", "단일-view 외관/치수 한계", "multi-view 또는 전문 route 자동 전환"),
                ("Blender bridge", "재질·렌더·assembly·USD", "느리고 version 의존", "asset library·headless regression"),
                ("DesignSpec 학습", "모델 출력 안전·검증 가능", "표현력 제한", "스키마 version과 허용 kind 확장"),
                ("DRF가 상태 소유", "재시도·감사·인증 일관성", "구성 요소 증가", "queue/tenant/observability"),
                ("내부 MySQL 먼저", "통제된 E2E와 schema 안정화", "사내 SSO 미연결", "corporate_db mapping 후 전환"),
                ("95% Wilson Gate", "작은 표본의 과대평가 차단", "많은 독립 데이터 필요", "범주별 ≥200 + blind review"),
                ("OpenUSD layer", "revision·assembly·simulation 연결", "운영 Nucleus 별도", "Asset Validator·WebRTC·권한"),
            ],
            [2200, 2640, 2200, 2320],
            font_size=7.8,
        )
        self.heading("14.1 현재 최우선 고도화 순서", 2)
        self.numbers(
            [
                "설비 외관 데이터셋과 전문 template을 늘리고 2D/3D component correspondence를 학습·검증한다.",
                "독립 VLM verifier와 사람이 분리된 고정 홀드아웃을 구축한다.",
                "OpenSCAD 구조를 Blender procedural asset/material library로 후처리한다.",
                "부분 재생성을 requirement group 단위로 정교화하고 실패 원인을 학습 데이터에 되먹임한다.",
                "실사진·도면·CAD의 권리 검토된 데이터로 domain gap을 줄인다.",
                "범주별 Wilson 95% 하한을 충족한 뒤에만 0.95 목표 달성을 선언한다.",
            ]
        )
        self.callout(
            "운영 판단",
            "현재 시스템은 설비 아이디어→구조화 요구→2D 비교→구조 3D/OpenUSD까지의 pre-CAD 도구로 사용할 수 있다. "
            "자동 제조 승인 도구로 사용하면 안 된다.",
            tone="warning",
        )

    def _checklists_and_sources(self) -> None:
        self.heading("15. 운영·승격 체크리스트", 1)
        self.heading("15.1 배포 전", 2)
        self.bullets(
            [
                "□ .env 비밀값이 Git/문서/이미지에 없음",
                "□ MySQL·storage·Qdrant 백업과 복구 drill 통과",
                "□ 100 service tests + 15 CAD tests 통과",
                "□ 직접 입력과 회의 음성 live E2E 통과",
                "□ UI login·history·viewer·download 확인",
                "□ provider model/revision/workflow SHA 기록",
                "□ 기능 Gate와 품질 Gate를 별도 승인",
                "□ 외관 0.95 미달이면 manufacturing 승격 금지",
            ]
        )
        self.heading("15.2 학습 승격", 2)
        self.bullets(
            [
                "□ 라이선스 allowlist와 training_allowed 확인",
                "□ split 누수·중복 이미지·경로 탈출·hash 검증 통과",
                "□ 10–50 step smoke와 checkpoint resume 확인",
                "□ 독립 test에서 범주별 ≥200",
                "□ Wilson 95% 하한 ≥0.95",
                "□ GeometryContract/native CAD/OpenUSD E2E 재실행",
                "□ 이전 승인 모델 대비 회귀 없음",
                "□ adapter·base revision·config·environment freeze를 묶음",
            ]
        )
        self.heading("15.3 제조 검토", 2)
        self.bullets(
            [
                "□ 실제 재질·공차·체결·표준부품 규격 확인",
                "□ 간섭·조립 순서·유지보수 공간 확인",
                "□ 구조 강도·열·전기·공압·안전 규격 검토",
                "□ STEP/B-Rep와 상세 feature tree를 CAD에서 보완",
                "□ 인증된 엔지니어와 제조 승인 담당자가 서명",
            ]
        )
        self.page_break()
        self.heading("부록 A. API 경로", 1)
        self.table(
            ["경로", "방법", "역할"],
            [
                ("/api/system-status", "GET", "provider·인증·native 상태"),
                ("/api/auth/config", "GET", "인증 모드"),
                ("/api/auth/login", "POST", "내부/외부 DB 로그인"),
                ("/api/auth/me", "GET", "현재 사용자"),
                ("/api/projects", "GET/POST", "목록/직접 입력 Project"),
                ("/api/meetings", "POST", "Meeting Project"),
                ("/api/projects/<id>", "GET", "Project 상세"),
                ("/api/projects/<id>/generate-3d", "POST", "선택 2D→3D"),
                ("/api/projects/<id>/meeting/chunks", "POST", "Audio chunk"),
                ("/api/projects/<id>/meeting/analyze", "POST", "회의 요구 구조화"),
                ("/api/projects/<id>/meeting/generate-2d", "POST", "회의→2D"),
                ("/api/projects/<id>/meeting/patch", "POST", "Revision patch"),
                ("/api/jobs/<uuid>", "GET", "비동기 상태"),
                ("/api/knowledge/ingest-file", "POST", "문서 RAG"),
                ("/api/knowledge/search", "POST", "과거 설계 검색"),
            ],
            [4300, 1500, 3560],
            font_size=8.0,
        )
        self.heading("부록 B. 실제 E2E 산출물", 1)
        self.table(
            ["상대 경로", "크기/내용"],
            [
                ("storage/projects/PRJ-E76ED59D94/concepts/concept-1.png", "805,265 bytes"),
                (".../result/structural/model_structural.glb", "54,884 bytes"),
                (".../result/structural/model_structural.stl", "594,635 bytes"),
                (".../result/structural/model.scad", "3,701 bytes"),
                (".../result/structural/geometry.json", "21,951 bytes"),
                (".../result/model.usda", "187,981 bytes"),
                (".../result/model.usdc", "28,639 bytes"),
                (".../result/openusd/manifest.json", "29,450 bytes"),
                (".../result/self_feedback_report.json", "score 0.4161 / target 0.95"),
                ("storage/e2e-evidence/.../ui-meeting-structured-result.png", "UI 증빙"),
            ],
            [6350, 3010],
            font_size=8.2,
        )
        self.heading("부록 C. 참고 문서와 외부 출처", 1)
        self.paragraph("저장소 기준 문서")
        self.bullets(
            [
                "README.md, docker-compose.yml, .env.example",
                "docs/ADVANCED_QUALITY_PROGRAM_KO.md",
                "docs/VALIDATION_AND_GRADE_KO.md",
                "docs/INTERNAL_MYSQL_TEST_KO.md",
                "training/cad-vlm/README.md",
                "training/cad-vlm/docs/TRAINING_AND_TUNING_MANUAL_KO.md",
                "training/cad-vlm/docs/MIGRATION_GUIDE_KO.md",
            ]
        )
        self.paragraph("음성 QA fixture 출처")
        sources = [
            (
                "Microsoft Speech 언어·한국어 음성 목록",
                "https://learn.microsoft.com/en-us/azure/ai-services/speech-service/language-support",
            ),
            ("edge-tts 프로젝트", "https://github.com/rany2/edge-tts"),
            ("Piper Korean KSS 모델", "https://huggingface.co/neurlang/piper-onnx-kss-korean"),
        ]
        for label, url in sources:
            paragraph = self.doc.add_paragraph(style="List Bullet")
            _add_hyperlink(paragraph, label, url)
        self.callout(
            "라이선스",
            "인터넷에서 받은 음성·CAD·이미지는 URL, revision, 다운로드 날짜, SHA-256, 라이선스 원문, "
            "training_allowed 승인을 기록한 뒤 사용한다. QA fixture 허용이 학습 데이터 허용을 뜻하지 않는다.",
            tone="warning",
        )
        self.heading("부록 D. 문서 변경 이력", 1)
        self.table(
            ["버전", "날짜", "내용"],
            [
                (
                    "1.0",
                    "2026-07-23",
                    "전체 파이프라인, 서버 이식, 라이브러리, 설정, CLI, 설계 이유, "
                    "PHP CAD 전처리 기반 학습, QA/E2E와 실제 검증 결과 통합",
                )
            ],
            [1300, 1800, 6260],
            font_size=8.8,
        )

    def save(self, output: Path) -> None:
        output.parent.mkdir(parents=True, exist_ok=True)
        core = self.doc.core_properties
        core.title = "Xconcep AI 전체 파이프라인·운영·학습·QA·E2E·서버 이식 매뉴얼"
        core.subject = "Verified technical operations manual"
        core.author = "Xconcep AI / Codex"
        core.keywords = "Xconcep, ComfyUI, FLUX, OpenSCAD, Blender, OpenUSD, CAD VLM, QA, E2E"
        core.comments = "Generated from the verified 2026-07-23 workspace state."
        self.doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    builder = ManualBuilder()
    builder.build()
    builder.save(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

from __future__ import annotations

import importlib.util
import json
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"C:\Users\audgh\OneDrive\문서\GPT Codex")
PROJECT = ROOT / "cad_ai_preprocessor_php"
BASE_BUILDER = ROOT / "manual_build" / "create_manual.py"
TEMPLATE = PROJECT / "docs_build" / "php_manual_content.md"
RESOLVED = PROJECT / "docs_build" / "php_manual_resolved.md"
OUT = ROOT / "DXF_STEP_AI_학습_데이터_전처리_PHP_풀스택_프로그램_매뉴얼_v2.1.docx"

spec = importlib.util.spec_from_file_location("base_manual_builder", BASE_BUILDER)
builder = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(builder)

builder.OUT = OUT


def resolve_content() -> None:
    report_path = PROJECT / "validation_reports" / "validation_10000_runs.json"
    values = {
        "VALIDATION_CYCLES": "검증 진행 중",
        "VALIDATION_TESTS": "검증 진행 중",
        "VALIDATION_PASS_RATE": "검증 진행 중",
        "VALIDATION_DURATION": "검증 진행 중",
        "VALIDATION_WORKERS": "검증 진행 중",
        "SOURCE_HASH": "검증 진행 중",
        "VALIDATION_P95": "검증 진행 중",
    }
    if report_path.exists():
        report = json.loads(report_path.read_text(encoding="utf-8-sig"))
        values.update(
            {
                "VALIDATION_CYCLES": f"{report['result']['successful_cycles']}/{report['scope']['iterations']}",
                "VALIDATION_TESTS": f"{report['result']['total_tests_run']:,}",
                "VALIDATION_PASS_RATE": f"{report['result']['pass_rate'] * 100:.2f}%",
                "VALIDATION_DURATION": f"{report['result']['total_duration_seconds']:.3f}초",
                "VALIDATION_WORKERS": str(report['scope']['parallel_workers']),
                "SOURCE_HASH": report["environment"]["source_sha256"],
                "VALIDATION_P95": f"{report['result']['p95_cycle_seconds']:.6f}초",
            }
        )
    text = TEMPLATE.read_text(encoding="utf-8")
    for key, value in values.items():
        text = text.replace("{{" + key + "}}", value)
    RESOLVED.write_text(text, encoding="utf-8")
    builder.CONTENT = RESOLVED


def header_footer(section):
    p = section.header.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(
        builder.Inches(6.5), builder.WD_TAB_ALIGNMENT.RIGHT
    )
    builder.font(p.add_run("CAD AI Dataset Studio · PHP"), size=8.5, color=builder.MUTED)
    builder.font(p.add_run("\t개발·사용·운영 매뉴얼 · V2.1"), size=8.5, color=builder.MUTED)
    builder.paragraph_border(p, "bottom", "D5DDE6", 4, 3)

    p = section.footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(
        builder.Inches(6.5), builder.WD_TAB_ALIGNMENT.RIGHT
    )
    builder.font(p.add_run("PHP 풀스택 · DXF·STEP 학습 데이터 전처리"), size=8, color=builder.MUTED)
    builder.font(p.add_run("\tPage "), size=8, color=builder.MUTED)
    builder.field(p, "PAGE")
    builder.font(p.add_run(" / "), size=8, color=builder.MUTED)
    builder.field(p, "NUMPAGES")


def cover(doc):
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    builder.font(
        p.add_run("PHP FULL-STACK · ENGINEERING DATA PIPELINE"),
        size=10,
        bold=True,
        color=builder.BLUE,
    )
    for text, size, after in (
        ("DXF·STEP 기반 AI 학습 데이터", 29, 4),
        ("PHP 풀스택 프로그램 개발·사용·운영 매뉴얼", 20, 14),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        builder.font(p.add_run(text), size=size, bold=True, color=builder.NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    builder.font(
        p.add_run("PHP 8.4 · HTML/CSS/JavaScript · SQLite/MySQL · 10,000회 검증"),
        size=13,
        color=builder.MUTED,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    builder.font(p.add_run("Version 2.1 · 2026-07-23"), size=11, bold=True, color=builder.DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    builder.font(
        p.add_run("기준: 업로드 개발 매뉴얼 · XconcepAI PHP 소스 · DXF·STEP 개발 방법"),
        size=9.5,
        color=builder.MUTED,
    )
    builder.callout(
        doc,
        "NOTE",
        "구현 기준",
        "웹·API·저장소·전처리·기준 학습은 PHP로 구현했다. "
        "정밀 STEP Assembly/B-Rep/Mesh가 필요한 생산 환경에서만 별도 OCP CAD Worker를 선택적으로 연동한다.",
    )
    doc.add_page_break()


def main():
    resolve_content()
    doc = builder.Document()
    builder.styles(doc)
    header_footer(doc.sections[0])
    nums = builder.numbering(doc)
    cover(doc)
    builder.parse(doc, nums)
    props = doc.core_properties
    props.title = "DXF·STEP 기반 AI 학습 데이터 전처리 PHP 풀스택 프로그램 매뉴얼"
    props.subject = "PHP 백엔드·웹 프론트엔드·DB·CAD 전처리·기준 모델·실제 CAD·10,000회 검증"
    props.author = "개발 조직"
    props.keywords = "PHP, DXF, STEP, AI 학습 데이터, 전처리, MySQL, Manifest"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def repair():
    resolve_content()
    builder.OUT = OUT
    builder.repair_lists()


if __name__ == "__main__":
    import sys

    repair() if "--repair-lists" in sys.argv else main()
